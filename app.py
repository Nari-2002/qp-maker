import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import re
import os

# --- Securely configure the Gemini API key ---
api_key = None
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
else:
    st.warning("`GEMINI_API_KEY` not found in `secrets.toml`. Please add it or enter it below.")
    api_key_input = st.text_input("Enter your Gemini API key:", type="password")
    if api_key_input:
        api_key = api_key_input
        genai.configure(api_key=api_key)

# --- App Title and Description ---
st.title("📝 AI-Powered Question Paper Generator")
st.markdown("Use Google Gemini to generate a custom multiple-choice question paper on any topic. You can select individual questions to build your final document.")

# --- Session State Initialization ---
if 'generated_questions' not in st.session_state:
    st.session_state.generated_questions = []
if 'selected_questions' not in st.session_state:
    st.session_state.selected_questions = []
if 'button_states' not in st.session_state:
    st.session_state.button_states = {}

# --- Functions to manage selected questions ---
def add_question_to_doc(question_data, q_index):
    question_text = question_data['question_text']
    if not any(q['question_text'] == question_text for q in st.session_state.selected_questions):
        question_data_with_index = question_data.copy()
        question_data_with_index['original_index'] = q_index
        st.session_state.selected_questions.append(question_data_with_index)
        
    st.session_state.button_states[q_index] = 'added'

def remove_question_from_doc(q_index):
    question_to_remove_text = st.session_state.generated_questions[q_index]['question_text']
    st.session_state.selected_questions = [
        q for q in st.session_state.selected_questions
        if q['question_text'] != question_to_remove_text
    ]
    
    st.session_state.button_states[q_index] = 'initial'

def remove_from_preview_and_reset_button(preview_index):
    if 0 <= preview_index < len(st.session_state.selected_questions):
        question_to_remove = st.session_state.selected_questions.pop(preview_index)
        original_index = question_to_remove.get('original_index')
        
        if original_index is not None and original_index in st.session_state.button_states:
            st.session_state.button_states[original_index] = 'initial'

def display_questions():
    """Renders the generated questions with 'Add/Undo' buttons based on button_states."""
    if st.session_state.generated_questions:
        st.header("2. Generated Questions")

        for i, q in enumerate(st.session_state.generated_questions):
            with st.container(border=True):
                st.markdown(f"**Question {i + 1}:** {q['question_text'].strip()}")
                
                options_lines = q['options_text'].strip().split('\n')
                for option_line in options_lines:
                    st.markdown(option_line.strip())
                
                if st.session_state.button_states.get(i) == 'added':
                    st.button("Undo", key=f"undo_button_{i}", on_click=remove_question_from_doc, args=(i,))
                else:
                    st.button("Add to Document", key=f"add_button_{i}", on_click=add_question_to_doc, args=(q, i))
                
                st.markdown("---")

# --- User Input Forms ---
st.header("1. Topics and Question Count")
topics_text = st.text_input(
    "Enter topics (comma-separated):",
    "Python Programming, Data Structures, Algorithms"
)

col1, col2, col3 = st.columns(3)
with col1:
    easy_questions = st.number_input("Number of Easy Questions", min_value=0, value=3)
with col2:
    medium_questions = st.number_input("Number of Medium Questions", min_value=0, value=3)
with col3:
    hard_questions = st.number_input("Number of Hard Questions", min_value=0, value=2)

# --- Question Generation Logic ---
def generate_questions_raw(topics, easy, medium, hard):
    """Generates questions from the API and returns the raw text."""
    if not api_key:
        return None
    try:
        prompt = f"""
        You are an expert question paper creator. Generate a question paper with multiple-choice questions (MCQs)
        based on the following topics: {topics}.
        The question paper should have a total of {easy + medium + hard} questions with the following difficulty distribution:
        - {easy} easy questions
        - {medium} medium questions
        - {hard} hard questions
        
        The questions should be a mix of all difficulty levels, not grouped by difficulty.
        Do not include difficulty levels like '(Easy)', '(Medium)', or '(Hard)' in the output.

        Each question should be a multiple-choice question with four options (A, B, C, D) and a single correct answer.
        The questions should be clear, concise, and directly related to the specified topics.
        Ensure the answers are accurate.

        Format the output clearly for easy parsing. Each question should follow this pattern:
        **[Question number]. [Question text]**
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        **Correct Answer: [Letter of the correct option]**

        Now, generate the question paper.
        """
        model = genai.GenerativeModel(model_name="gemini-2.5-flash-preview-05-20")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"An error occurred while generating questions: {e}")
        return None

def parse_questions(raw_text):
    """Parses raw text into a list of dictionaries for a more robust handling."""
    questions_list = []
    
    # Split the raw text by the question number header
    question_blocks = re.split(r'\*\*(\d+\.\s.*?)\*\*', raw_text, flags=re.DOTALL)
    
    # The first element will be empty, so we skip it
    for i in range(1, len(question_blocks), 2):
        question_text_raw = question_blocks[i]
        content_block = question_blocks[i+1]
        
        # Now, parse the content block for options and correct answer
        options_and_answer = re.search(r'(A\).+?)\s*\*\*Correct Answer:\s*([A-D])\s*', content_block, re.DOTALL)
        
        if options_and_answer:
            options_block = options_and_answer.group(1)
            correct_answer = options_and_answer.group(2)
            
            # Split the options block into individual options
            options_list = re.findall(r'([A-D]\).+?)(?=\s*[A-D]\)|\s*$)', options_block, re.DOTALL)
            options_text = "\n".join([opt.strip() for opt in options_list])
            
            questions_list.append({
                "question_text": question_text_raw.strip(),
                "options_text": options_text,
                "correct_answer": correct_answer.strip()
            })
    
    return questions_list


# --- Main App Flow ---
if st.button("Generate Question Paper", key="generate_button"):
    if not api_key:
        st.warning("Please enter your Gemini API key to proceed.")
    elif not topics_text:
        st.warning("Please enter at least one topic.")
    else:
        with st.spinner("Generating questions... This may take a moment."):
            raw_response = generate_questions_raw(topics_text, easy_questions, medium_questions, hard_questions)
            if raw_response:
                st.session_state.generated_questions = parse_questions(raw_response)
                st.session_state.button_states = {}
                st.success("Questions generated successfully! You can now select new questions or keep your existing ones.")

display_questions()

# --- Preview and Download Section ---
if st.session_state.selected_questions:
    st.header("3. Preview and Download")
    st.info(f"{len(st.session_state.selected_questions)} questions selected.")
    
    st.subheader("Selected Questions Preview")
    for i, q in enumerate(st.session_state.selected_questions):
        with st.expander(f"Question {i + 1}: {q['question_text'].split('. ', 1)[-1]}..."):
            st.markdown(f"**Question:** {q['question_text'].split('. ', 1)[-1]}")
            st.markdown(q['options_text'])
            st.markdown(f"**Correct Answer:** {q['correct_answer']}")
            
            st.button(
                "Remove from Document", 
                key=f"remove_preview_{i}", 
                on_click=remove_from_preview_and_reset_button, 
                args=(i,)
            )

    def create_word_document_from_selection(selected_questions):
        document = Document()
        document.add_heading('Generated Question Paper', 0)
        
        for i, q in enumerate(selected_questions):
            question_para = document.add_paragraph()
            question_text_clean = q['question_text'].split('. ', 1)[-1].strip()
            question_para.add_run(f"{i + 1}. {question_text_clean}").bold = True

            options_lines = q['options_text'].strip().split('\n')
            for option_line in options_lines:
                if option_line.strip():
                    document.add_paragraph(option_line.strip())

            answer_para = document.add_paragraph()
            answer_para.add_run(f"Correct Answer: {q['correct_answer']}").italic = True
        
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return file_stream

    docx_file = create_word_document_from_selection(st.session_state.selected_questions)
    st.download_button(
        label="Download Final Question Paper (.docx)",
        data=docx_file,
        file_name="question_paper.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )