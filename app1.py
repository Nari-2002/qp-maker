import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
import re

# --- Securely configure the Gemini API key ---
# Streamlit secrets is the recommended way for production apps.
api_key = st.secrets.get("GEMINI_API_KEY")

if not api_key:
    st.warning("`GEMINI_API_KEY` not found in `secrets.toml`. Please add it or enter it below.")
    api_key = st.text_input("Enter your Gemini API key:", type="password")

if api_key:
    genai.configure(api_key=api_key)

# --- App Title and Description ---
st.set_page_config(page_title="AI-Powered Question Paper Generator", layout="wide")
st.title("📝 AI-Powered Question Paper Generator")
st.markdown("Use Google Gemini to generate a custom multiple-choice question paper on any topic. You can select individual questions to build your final document.")

# --- Session State Initialization ---
if 'all_generated_questions' not in st.session_state:
    st.session_state.all_generated_questions = []
if 'current_page' not in st.session_state:
    st.session_state.current_page = 0
if 'selected_questions' not in st.session_state:
    st.session_state.selected_questions = []
if 'button_states' not in st.session_state:
    st.session_state.button_states = {}

# --- Functions to manage selected questions and pagination ---
def add_question_to_doc(question_data, page_index, q_index):
    """Adds a question to the selected list and updates button state."""
    question_text = question_data['question_text']
    
    if not any(q['question_text'] == question_text for q in st.session_state.selected_questions):
        question_data_with_meta = question_data.copy()
        question_data_with_meta['page_index'] = page_index
        question_data_with_meta['original_index'] = q_index
        st.session_state.selected_questions.append(question_data_with_meta)
    
    key = (page_index, q_index)
    st.session_state.button_states[key] = 'added'

def remove_question_from_doc(page_index, q_index):
    """Removes a question from the selected list and updates button state."""
    question_to_remove_text = st.session_state.all_generated_questions[page_index][q_index]['question_text']
    
    st.session_state.selected_questions = [
        q for q in st.session_state.selected_questions
        if q['question_text'] != question_to_remove_text
    ]
    
    key = (page_index, q_index)
    st.session_state.button_states[key] = 'initial'

def remove_from_preview_and_reset_button(preview_index):
    """Removes a question from the preview and resets its corresponding button state."""
    if 0 <= preview_index < len(st.session_state.selected_questions):
        question_to_remove = st.session_state.selected_questions.pop(preview_index)
        page_index = question_to_remove.get('page_index')
        original_index = question_to_remove.get('original_index')
        
        if page_index is not None and original_index is not None:
            key = (page_index, original_index)
            if key in st.session_state.button_states:
                st.session_state.button_states[key] = 'initial'

def display_questions():
    """Renders the generated questions for the current page with pagination controls."""
    if not st.session_state.all_generated_questions:
        st.info("Click 'Generate New Questions' to create your first set of questions.")
        return

    st.header("2. Generated Questions")
    
    # Pagination Controls
    total_pages = len(st.session_state.all_generated_questions)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.session_state.current_page > 0:
            if st.button("Previous Page", key="prev_page"):
                st.session_state.current_page -= 1
                st.experimental_rerun()
    with col2:
        st.markdown(f"<h5 style='text-align: center;'>Page {st.session_state.current_page + 1} of {total_pages}</h5>", unsafe_allow_html=True)
    with col3:
        if st.session_state.current_page < total_pages - 1:
            if st.button("Next Page", key="next_page"):
                st.session_state.current_page += 1
                st.experimental_rerun()
    
    st.markdown("---")
    
    current_page_questions = st.session_state.all_generated_questions[st.session_state.current_page]
    
    for i, q in enumerate(current_page_questions):
        with st.container(border=True):
            st.markdown(f"**Question {i + 1}:** {q['question_text'].strip()}")
            
            # Use columns for options for a cleaner layout
            options_lines = q['options_text'].strip().split('\n')
            num_options = len(options_lines)
            cols = st.columns(num_options)
            for col, option_line in zip(cols, options_lines):
                col.markdown(option_line.strip())
            
            key = (st.session_state.current_page, i)
            
            if st.session_state.button_states.get(key) == 'added':
                if st.button("Undo", key=f"undo_{key}"):
                    remove_question_from_doc(st.session_state.current_page, i)
                    st.experimental_rerun()
            else:
                if st.button("Add to Document", key=f"add_{key}"):
                    add_question_to_doc(q, st.session_state.current_page, i)
                    st.experimental_rerun()
            
            st.markdown("---")

# --- User Input Forms ---
st.sidebar.header("1. Topics and Question Count")
topics_text = st.sidebar.text_input(
    "Enter topics (comma-separated):",
    "Python Programming, Data Structures, Algorithms"
)

col1, col2, col3 = st.sidebar.columns(3)
with col1:
    easy_questions = st.number_input("Easy", min_value=0, value=3)
with col2:
    medium_questions = st.number_input("Medium", min_value=0, value=3)
with col3:
    hard_questions = st.number_input("Hard", min_value=0, value=2)

# --- Question Generation Logic ---
def generate_questions_raw(topics, easy, medium, hard):
    """Generates questions from the API and returns the raw text."""
    if not api_key:
        st.warning("Please enter your Gemini API key to proceed.")
        return None
    
    try:
        total_questions = easy + medium + hard
        prompt = f"""
        You are an expert question paper creator. Generate a question paper with multiple-choice questions (MCQs)
        based on the following topics: {topics}.
        The question paper should have a total of {total_questions} questions with the following difficulty distribution:
        - {easy} easy questions
        - {medium} medium questions
        - {hard} hard questions
        
        The questions should be a mix of all difficulty levels, not grouped by difficulty.
        Do not include difficulty levels like '(Easy)', '(Medium)', or '(Hard)' in the output.

        Each question should be a multiple-choice question with four options (A, B, C, D) and a single correct answer.
        The questions should be clear, concise, and directly related to the specified topics.
        Ensure the answers are accurate.

        Format the output clearly for easy parsing. Each question should follow this pattern:
        **[Question number]. [Question text]**
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        **Correct Answer: [Letter of the correct option]**

        Now, generate the question paper.
        """
        model = genai.GenerativeModel(model_name="gemini-1.5-flash")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"An error occurred while generating questions: {e}")
        return None

def parse_questions(raw_text):
    """Parses raw text into a list of dictionaries. Updated regex for robustness."""
    questions_list = []
    
    # Updated regex to handle different spacing and formatting issues
    question_pattern = re.compile(
        r'\*\*(\d+\.\s.*?)\*\*\s*'  # Matches question number and text
        r'(A\).+?)\s*'             # Matches option A
        r'(B\).+?)\s*'             # Matches option B
        r'(C\).+?)\s*'             # Matches option C
        r'(D\).+?)\s*'             # Matches option D
        r'\*\*Correct Answer:\s*([A-D])\s*', # Matches the correct answer
        re.DOTALL | re.IGNORECASE
    )
    
    matches = question_pattern.findall(raw_text)
    
    for match in matches:
        question_text_raw = match[0].strip()
        options = [match[1], match[2], match[3], match[4]]
        options_text = "\n".join(opt.strip() for opt in options)
        correct_answer = match[5].strip()
        
        questions_list.append({
            "question_text": question_text_raw.strip(),
            "options_text": options_text,
            "correct_answer": correct_answer
        })
        
    return questions_list

# --- Main App Flow ---
if st.sidebar.button("Generate New Questions", key="generate_button"):
    if not api_key:
        st.warning("Please enter your Gemini API key to proceed.")
    elif not topics_text:
        st.warning("Please enter at least one topic.")
    else:
        with st.spinner("Generating questions... This may take a moment."):
            raw_response = generate_questions_raw(topics_text, easy_questions, medium_questions, hard_questions)
            if raw_response:
                new_questions_page = parse_questions(raw_response)
                if new_questions_page:
                    st.session_state.all_generated_questions.append(new_questions_page)
                    st.session_state.current_page = len(st.session_state.all_generated_questions) - 1
                    st.success("New questions generated successfully! Use the navigation to view them.")
                    st.experimental_rerun()
                else:
                    st.error("Failed to parse questions. The generated text format may have changed. Please try again.")
                    st.text_area("Raw Response from API (for debugging)", raw_response, height=300)

display_questions()

# --- Preview and Download Section ---
st.header("3. Download Question Paper")

logo_file = st.file_uploader("Upload School/Company Logo (optional):", type=['png', 'jpg', 'jpeg'])

if st.session_state.selected_questions:
    st.info(f"{len(st.session_state.selected_questions)} questions selected.")
    
    with st.expander("Preview and Reorder Selected Questions"):
        st.markdown("Drag and drop questions to reorder them.")
        
        # Display selected questions with drag-and-drop capability
        st.session_state.selected_questions = st.data_editor(
            st.session_state.selected_questions,
            column_order=["question_text"],
            column_config={
                "question_text": st.column_config.TextColumn("Question", width="medium"),
                "options_text": None,
                "correct_answer": None,
                "page_index": None,
                "original_index": None
            },
            hide_index=True,
            num_rows="dynamic",
            use_container_width=True
        )

    # --- Create and Download Word Document ---
    def create_word_document_from_selection(selected_questions, logo_data):
        document = Document()
        
        # Header
        header = document.sections[0].header
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        
        left_cell = header_table.cell(0, 0)
        left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if logo_data:
            logo_stream = BytesIO(logo_data.getvalue())
            left_cell.paragraphs[0].add_run().add_picture(logo_stream, width=Inches(1.5))
        
        right_cell = header_table.cell(0, 1)
        right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.paragraphs[0].line_spacing = Pt(12)
        
        right_cell.add_paragraph('Lakshmipuraim Main Road').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.add_paragraph('Opp. Sri Patibandla Sitaramaiah High School').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.add_paragraph('Guntur, Andhra Pradesh.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.add_paragraph('Ph: +91 8247729604, 8179423238').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.add_paragraph('Email: learn@faangtechlab.com').alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_cell.add_paragraph('Web: www.faangtechlab.com').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Main content
        num_questions = len(selected_questions)
        questions_per_column = (num_questions + 1) // 2
        
        # Use a table for a two-column layout
        main_table = document.add_table(rows=questions_per_column, cols=2)
        main_table.autofit = False
        main_table.allow_autofit = False
        
        for cell in main_table.columns[0].cells:
            cell.width = Inches(3.0)
        for cell in main_table.columns[1].cells:
            cell.width = Inches(3.0)
        
        def set_cell_border(cell, **kwargs):
            tc = cell._element.get_or_add_tcPr()
            for key, value in kwargs.items():
                tag = OxmlElement('w:{}'.format(key))
                tag.set(ns.qn('w:val'), value)
                tc.append(tag)
        
        for row in main_table.rows:
            for cell in row.cells:
                set_cell_border(cell, top='nil', left='nil', bottom='nil', right='nil')

        current_row = 0
        current_col = 0
        
        for q_index, q in enumerate(selected_questions):
            cell = main_table.cell(current_row, current_col)

            question_para = cell.add_paragraph()
            question_text_clean = q['question_text'].strip()
            run = question_para.add_run(f"{q_index + 1}. {question_text_clean}")
            run.font.bold = True
            
            options_lines = q['options_text'].strip().split('\n')
            for option_line in options_lines:
                option_para = cell.add_paragraph(option_line.strip())
                option_para.paragraph_format.left_indent = Inches(0.25)
            
            cell.add_paragraph('')

            if current_col == 0 and q_index < questions_per_column - 1:
                current_col = 1
            else:
                current_col = 0
                current_row += 1
        
        # Create a separate document for answers
        answers_doc = Document()
        answers_doc.add_heading("Answer Key", 1)
        for q_index, q in enumerate(selected_questions):
            answers_doc.add_paragraph(f"{q_index + 1}. {q['correct_answer']}")
        
        # Save both documents to streams
        question_paper_stream = BytesIO()
        document.save(question_paper_stream)
        question_paper_stream.seek(0)
        
        answers_key_stream = BytesIO()
        answers_doc.save(answers_key_stream)
        answers_key_stream.seek(0)

        return question_paper_stream, answers_key_stream

    # Create the documents and download buttons
    docx_file, answers_file = create_word_document_from_selection(st.session_state.selected_questions, logo_file)
    
    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button(
            label="Download Question Paper (.docx)",
            data=docx_file,
            file_name="question_paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_dl2:
        st.download_button(
            label="Download Answer Key (.docx)",
            data=answers_file,
            file_name="answer_key.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("No questions have been added to the document yet.")