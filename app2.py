import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from io import BytesIO
from datetime import datetime
import re

# --- Securely configure the Gemini API key ---
api_key = st.secrets.get("GEMINI_API_KEY")
if not api_key:
    st.warning("`GEMINI_API_KEY` not found in `secrets.toml`. Please add it or enter it below.")
    api_key = st.text_input("Enter your Gemini API key:", type="password")
if api_key:
    genai.configure(api_key=api_key)
# --- Page Configuration ---
st.set_page_config(
    page_title="My App",
    page_icon="📝" # Use an emoji here
)
# --- Page Setup ---
st.set_page_config(page_title="Question Paper Generator", layout="wide")
st.title("📝 Question Paper Generator")
st.markdown("""
Generate custom multiple-choice question papers using Google Gemini. Select, reorder, and download questions with an answer key.
""")

# --- Session Initialization ---
if 'all_generated_questions' not in st.session_state:
    st.session_state.all_generated_questions = []
if 'current_page' not in st.session_state:
    st.session_state.current_page = 0
if 'selected_questions' not in st.session_state:
    st.session_state.selected_questions = []
if 'button_states' not in st.session_state:
    st.session_state.button_states = {}
if 'show_answers_in_doc' not in st.session_state:
    st.session_state.show_answers_in_doc = False

# --- Sidebar: User Inputs ---
st.sidebar.header("1. Topics and Settings")
topics_text = st.sidebar.text_input("Enter topics (comma-separated):", "Python, Data Structures")
col1, col2, col3 = st.sidebar.columns(3)
easy = col1.number_input("Easy", min_value=0, value=3)
medium = col2.number_input("Medium", min_value=0, value=3)
hard = col3.number_input("Hard", min_value=0, value=2)

st.sidebar.markdown("---")
st.sidebar.toggle("Show correct answers in question paper?", key="show_answers_in_doc")

# --- Generation Logic ---
def generate_questions():
    """Generates questions from Gemini and stores them in session state."""
    if not api_key:
        st.warning("Please enter your Gemini API key to proceed.")
        return
    if not topics_text:
        st.warning("Please enter at least one topic.")
        return

    with st.spinner("Generating questions..."):
        total = easy + medium + hard
        if total == 0:
            st.warning("Please specify at least one question.")
            return

        prompt = f"""
        You are an expert MCQ creator. Generate {total} multiple-choice questions on the topic(s): {topics_text}.
        The questions should have the following difficulty distribution:
        - {easy} easy questions
        - {medium} medium questions
        - {hard} hard questions
        Do not include the difficulty level in the output.

        Each question must be a multiple-choice question with four options (A, B, C, D) and a single correct answer.
        Ensure the questions are clear, concise, and directly related to the specified topics.

        Format the output precisely as follows for each question:
        **1. [Question text]**
        A) [Option A]
        B) [Option B]
        C) [Option C]
        D) [Option D]
        **Correct Answer: [Letter of the correct option]**
        """
        model = genai.GenerativeModel(model_name="gemini-1.5-flash")
        try:
            response = model.generate_content(prompt)
            raw_text = response.text
            
            # The regex is now more robust to handle slight variations in output
            pattern = re.compile(
                r'\*\*(\d+\. .*?)\*\*\s*'
                r'(A\).+?)\s*'
                r'(B\).+?)\s*'
                r'(C\).+?)\s*'
                r'(D\).+?)\s*'
                r'\*\*Correct Answer:\s*([A-D])\s*',
                re.DOTALL | re.IGNORECASE
            )
            matches = pattern.findall(raw_text)
            
            page = []
            if not matches and "```" in raw_text:
                # Handle cases where the model wraps output in markdown code blocks
                clean_text = raw_text.replace("```", "").strip()
                matches = pattern.findall(clean_text)

            for match in matches:
                qtext = match[0].strip()
                options = "\n".join(m.strip() for m in match[1:5])
                answer = match[5].strip()
                page.append({"question_text": qtext, "options_text": options, "correct_answer": answer})
            
            if page:
                st.session_state.all_generated_questions.append(page)
                st.session_state.current_page = len(st.session_state.all_generated_questions) - 1
                st.success("Questions generated!")
                st.rerun()
            else:
                st.error("Failed to parse questions. The generated text format may have changed. Please try again.")
                with st.expander("Show raw output for debugging"):
                    st.text_area("Raw Output:", raw_text, height=300)
        except Exception as e:
            st.error(f"Error: {e}")

if st.sidebar.button("🔄 Generate New Questions"):
    generate_questions()

# --- Display Sample If No Questions ---
if not st.session_state.all_generated_questions:
    st.info("No questions generated yet. Try entering a topic and click 'Generate New Questions'.")
    with st.expander("📌 Sample Question Format"):
        st.markdown("""
        **1. What is Python?**
        A) A type of snake
        B) A high-level programming language
        C) A famous book
        D) A type of car

        **Correct Answer: B**
        """)

# --- Question Display and Pagination ---
def display_questions():
    """Displays generated questions with pagination controls."""
    total_pages = len(st.session_state.all_generated_questions)
    if total_pages == 0:
        return

    st.subheader("2. Generated Questions")
    
    # Pagination controls at the top of the section
    col_nav1, col_nav2, col_nav3 = st.columns([1, 2, 1])
    with col_nav1:
        if st.session_state.current_page > 0:
            if st.button("⬅ Previous Page"):
                st.session_state.current_page -= 1
                st.rerun()
    with col_nav2:
        st.markdown(f"<h5 style='text-align: center;'>Page {st.session_state.current_page + 1} of {total_pages}</h5>", unsafe_allow_html=True)
    with col_nav3:
        if st.session_state.current_page < total_pages - 1:
            if st.button("Next Page ➡"):
                st.session_state.current_page += 1
                st.rerun()

    st.markdown("---")
    
    page_qs = st.session_state.all_generated_questions[st.session_state.current_page]

    for i, q in enumerate(page_qs):
        key = (st.session_state.current_page, i)
        with st.container(border=True):
            st.markdown(f"**{i+1}. {q['question_text']}**")
            for opt in q['options_text'].split('\n'):
                st.markdown(f"{opt}")
            
            # Use columns for the Add/Undo buttons for a cleaner layout
            btn_col, _ = st.columns([1, 5])
            with btn_col:
                if st.session_state.button_states.get(key) == 'added':
                    if st.button("❌ Undo", key=f"undo_{key}"):
                        q_text = q['question_text']
                        st.session_state.selected_questions = [s for s in st.session_state.selected_questions if s['question_text'] != q_text]
                        st.session_state.button_states[key] = 'initial'
                        st.rerun()
                else:
                    if st.button("➕ Add", key=f"add_{key}"):
                        # Ensure the question isn't already added
                        if not any(s['question_text'] == q['question_text'] for s in st.session_state.selected_questions):
                            q_with_meta = q.copy()
                            q_with_meta['page_index'] = st.session_state.current_page
                            q_with_meta['original_index'] = i
                            st.session_state.selected_questions.append(q_with_meta)
                            st.session_state.button_states[key] = 'added'
                            st.rerun()
            st.markdown("---")

# --- Question Preview & Reordering ---
if st.session_state.selected_questions:
    st.subheader("🧾 Selected Questions Preview")
    st.info(f"Total Questions Selected: {len(st.session_state.selected_questions)}")
    st.session_state.selected_questions = st.data_editor(
        st.session_state.selected_questions,
        column_order=["question_text"],
        column_config={"question_text": st.column_config.TextColumn("Question", help="Drag and drop to reorder.")},
        hide_index=True,
        use_container_width=True
    )

# --- Exam Header Info ---
st.header("3. Download Question Paper")
st.text_input("Exam Name:", value="Mid-Term Exam", key="exam_name")
st.text_input("Duration (mins):", value="90", key="exam_duration")
st.text_input("Date:", value=str(datetime.today().date()), key="exam_date")
logo_file = st.file_uploader("Upload Logo:", type=['png', 'jpg', 'jpeg'])

# --- Clear Button ---
if st.button("🗑️ Clear All Selected Questions"):
    st.session_state.selected_questions = []
    st.session_state.button_states = {}
    st.rerun() # Use rerun to reflect the change immediately

# --- Create & Download DOCX ---
def create_docs(questions, show_answers, logo):
    """Creates the main question paper and a separate answer key document."""
    # Main Question Paper Document
    doc = Document()
    section = doc.sections[0]
    header = section.header
    header_table = header.add_table(1, 2, Inches(6.5))
    
    left_cell = header_table.cell(0, 0)
    left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo:
        img_stream = BytesIO(logo.getvalue())
        left_cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(1.5))
    
    right_cell = header_table.cell(0, 1)
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_cell.paragraphs[0].line_spacing = Pt(12)
    right_cell.text = f"{st.session_state.exam_name}\nDate: {st.session_state.exam_date}\nDuration: {st.session_state.exam_duration} mins"

    for i, q in enumerate(questions):
        doc.add_paragraph(f"{i+1}. {q['question_text']}", style='List Number')
        for opt in q['options_text'].split('\n'):
            doc.add_paragraph(opt, style='List Bullet')
        if show_answers:
            doc.add_paragraph(f"Correct Answer: {q['correct_answer']}", style='Normal').font.bold = True
            
    main_stream = BytesIO()
    doc.save(main_stream)
    main_stream.seek(0)
    
    # Answer Key Document
    key_doc = Document()
    key_doc.add_heading("Answer Key", 1)
    for i, q in enumerate(questions):
        key_doc.add_paragraph(f"{i+1}. {q['correct_answer']}")

    key_stream = BytesIO()
    key_doc.save(key_stream)
    key_stream.seek(0)
    
    return main_stream, key_stream

if st.session_state.selected_questions:
    docx_file, key_file = create_docs(st.session_state.selected_questions, st.session_state.show_answers_in_doc, logo_file)
    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        st.download_button(
            "⬇️ Download Question Paper", 
            docx_file, 
            file_name="question_paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_dl2:
        st.download_button(
            "⬇️ Download Answer Key", 
            key_file, 
            file_name="answer_key.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- Call the display function to render the UI ---
display_questions()